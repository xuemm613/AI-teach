"""智能备课接口。"""
import logging
from datetime import datetime
from io import BytesIO
from typing import Optional
from urllib.parse import quote

import docx
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from sqlalchemy import desc, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import require_roles
from app.api.subject_utils import is_valid_subject, teacher_subjects
from app.core.database import get_db
from app.models.models import LessonPlan, User
from app.schemas.common import ok
from app.schemas.lesson import LessonGenerateRequest, LessonPlanUpdate
from app.services.lesson_service import lesson_service

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/lesson", tags=["智能备课"])


def plan_to_dict(plan: LessonPlan) -> dict:
    return {
        "id": plan.id,
        "teacher_id": plan.teacher_id,
        "grade": plan.grade,
        "subject": plan.subject,
        "chapter": plan.chapter,
        "teaching_objectives": plan.teaching_objectives,
        "content": plan.content,
        "status": plan.status,
        "created_at": plan.created_at.isoformat() if plan.created_at else None,
    }


@router.post("/generate")
async def generate_lesson(
    payload: LessonGenerateRequest,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    # 学科必须为系统规定科目
    if not is_valid_subject(payload.subject):
        return fail("学科必须为系统规定科目：语文、数学、英语、物理、化学、生物、政治、历史、地理、体育、音乐、美术、劳动")
    # 教师只能备课自己负责的科目
    if user.role == "teacher":
        subs = await teacher_subjects(db, user)
        if payload.subject not in subs:
            return fail(f"只能备课您负责的科目：{', '.join(sorted(subs)) or '未分配科目'}")

    plan = await lesson_service.generate(
        db=db,
        teacher_id=user.id,
        grade=payload.grade,
        subject=payload.subject,
        chapter=payload.chapter,
        teaching_objectives=payload.teaching_objectives,
    )
    return ok(plan_to_dict(plan), message="教案生成成功")


@router.get("/plans")
async def list_plans(
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
    start: Optional[str] = Query(None, description="开始日期 YYYY-MM-DD"),
    end: Optional[str] = Query(None, description="结束日期 YYYY-MM-DD"),
    keyword: Optional[str] = Query(None, description="按章节关键词检索"),
):
    stmt = select(LessonPlan)
    if user.role == "teacher":
        stmt = stmt.where(LessonPlan.teacher_id == user.id)
    if start:
        stmt = stmt.where(LessonPlan.created_at >= datetime.fromisoformat(start))
    if end:
        stmt = stmt.where(LessonPlan.created_at <= datetime.fromisoformat(end) + __import__("datetime").timedelta(days=1))
    if keyword:
        stmt = stmt.where(LessonPlan.chapter.ilike(f"%{keyword}%"))
    result = await db.execute(stmt.order_by(desc(LessonPlan.created_at)))
    plans = result.scalars().all()
    return ok([plan_to_dict(p) for p in plans])


@router.put("/plans/{plan_id}")
async def update_plan(
    plan_id: int,
    payload: LessonPlanUpdate,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    """手动编辑/微调教案内容。"""
    plan = await db.get(LessonPlan, plan_id)
    if plan is None:
        raise HTTPException(status_code=404, detail="教案不存在")
    if user.role == "teacher" and plan.teacher_id != user.id:
        raise HTTPException(status_code=403, detail="无权编辑该教案")
    plan.content = payload.content
    plan.status = "edited"
    await db.commit()
    await db.refresh(plan)
    return ok(plan_to_dict(plan), message="教案已保存")


@router.get("/plans/{plan_id}/export")
async def export_plan_word(
    plan_id: int,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    """导出教案为 Word (.docx)。"""
    plan = await db.get(LessonPlan, plan_id)
    if plan is None:
        raise HTTPException(status_code=404, detail="教案不存在")
    if user.role == "teacher" and plan.teacher_id != user.id:
        raise HTTPException(status_code=403, detail="无权导出该教案")

    content = plan.content or {}
    doc = docx.Document()
    doc.add_heading(f"{plan.subject} · {plan.chapter} 教案", 0)

    doc.add_heading("一、教学目标", level=1)
    for obj in content.get("teaching_objectives") or []:
        doc.add_paragraph(str(obj), style="List Bullet")

    doc.add_heading("二、课堂导入", level=1)
    doc.add_paragraph(content.get("introduction") or "")

    doc.add_heading("三、讲授提纲", level=1)
    for i, item in enumerate(content.get("outline") or [], start=1):
        doc.add_paragraph(f"{i}. {item}")

    doc.add_heading("四、互动问题", level=1)
    for i, q in enumerate(content.get("interactive_questions") or [], start=1):
        doc.add_paragraph(f"问{i}：{q.get('question', '')}")
        doc.add_paragraph(f"答：{q.get('answer', '')}")

    doc.add_heading("五、板书设计", level=1)
    doc.add_paragraph(content.get("board_design") or "")

    doc.add_heading("六、分层练习", level=1)
    layers = content.get("layered_exercises") or {}
    for label, key in (("基础题", "basic"), ("提高题", "medium"), ("拓展题", "advanced")):
        doc.add_paragraph(label + "：")
        for ex in layers.get(key) or []:
            doc.add_paragraph(str(ex), style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = quote(f"{plan.subject}_{plan.chapter}_教案.docx")
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


@router.delete("/plans/{plan_id}")
async def delete_plan(
    plan_id: int,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    plan = await db.get(LessonPlan, plan_id)
    if plan is None:
        raise HTTPException(status_code=404, detail="教案不存在")
    if user.role == "teacher" and plan.teacher_id != user.id:
        raise HTTPException(status_code=403, detail="无权删除该教案")
    await db.delete(plan)
    await db.commit()
    return ok(None, message="删除成功")
