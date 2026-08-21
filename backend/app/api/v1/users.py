"""用户接口：个人信息、学生端（看板/问答记录/练习/错题本/时间线/个人中心）、
教师端（班级/学情/留言）、用户管理。"""
import logging
from collections import Counter, defaultdict
from datetime import date, datetime, timedelta
from io import BytesIO
from typing import List, Optional
from urllib.parse import quote

import docx
from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from fastapi.responses import Response
from sqlalchemy import case, delete, desc, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_user, require_roles
from app.api.subject_utils import is_valid_subject
from app.core.config import settings
from app.core.database import get_db
from app.core.security import hash_password
from app.models.models import (
    AgentTask,
    Class,
    ClassSchedule,
    ClassScheduleTeacher,
    ClassStudent,
    ChatMessage,
    ChatSession,
    Course,
    Exercise,
    LearningRecord,
    LessonPlan,
    Student,
    StudentMessage,
    Teacher,
    User,
    WrongBook,
)
from app.schemas.common import fail, ok
from app.schemas.user import (
    AdminUserCreate,
    AdminUserUpdate,
    StudentMessageCreate,
    SubmitRecordRequest,
    UserUpdate,
    WrongBookAddRequest,
)
from app.core.llm import llm_client, parse_json
from app.services.kp_similarity import ExerciseFeature, KP_SYNONYMS, canonical_kps, expanded_kps, select_similar
from app.utils.prompt_templates import SIMILAR_EXERCISE_PROMPT

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/users", tags=["用户"])

AVATAR_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}


def exercise_to_dict(ex: Exercise) -> dict:
    return {
        "id": ex.id,
        "course_id": ex.course_id,
        "chapter": ex.chapter,
        "type": ex.type,
        "content": ex.content,
        "options": ex.options or [],
        "answer": ex.answer,
        "analysis": ex.analysis,
        "difficulty": ex.difficulty,
        "knowledge_points": ex.knowledge_points or [],
    }


def user_to_dict(user: User, db_student: Optional[Student] = None, db_teacher: Optional[Teacher] = None) -> dict:
    return {
        "id": user.id,
        "username": user.username,
        "full_name": user.full_name,
        "email": user.email,
        "role": user.role,
        "avatar": user.avatar,
        "is_active": user.is_active,
        "created_at": user.created_at.isoformat() if user.created_at else None,
        "student_no": db_student.student_no if db_student else None,
        "grade": db_student.grade if db_student else None,
        "title": db_teacher.title if db_teacher else None,
        "department": db_teacher.department if db_teacher else None,
        "employee_no": db_teacher.employee_no if db_teacher else None,
        "subjects": list(db_teacher.subjects or []) if db_teacher else [],
    }


async def _load_profile(db: AsyncSession, user: User) -> tuple:
    student = teacher = None
    if user.role == "student":
        result = await db.execute(select(Student).where(Student.user_id == user.id))
        student = result.scalar_one_or_none()
        if student is None:
            student = Student(user_id=user.id, student_no=f"S{user.id:06d}")
            db.add(student)
            await db.commit()
            await db.refresh(student)
    elif user.role == "teacher":
        result = await db.execute(select(Teacher).where(Teacher.user_id == user.id))
        teacher = result.scalar_one_or_none()
        if teacher is None:
            teacher = Teacher(user_id=user.id)
            db.add(teacher)
            await db.commit()
            await db.refresh(teacher)
    return student, teacher


async def _get_student(db: AsyncSession, user: User) -> Student:
    result = await db.execute(select(Student).where(Student.user_id == user.id))
    student = result.scalar_one_or_none()
    if student is None:
        raise HTTPException(status_code=400, detail="当前账号非学生账号")
    return student


async def _get_teacher(db: AsyncSession, user: User) -> Teacher:
    result = await db.execute(select(Teacher).where(Teacher.user_id == user.id))
    teacher = result.scalar_one_or_none()
    if teacher is None:
        raise HTTPException(status_code=400, detail="当前账号非教师账号")
    return teacher


async def _knowledge_mastery(
    db: AsyncSession, student_id: int, limit: int = 1000, subject: Optional[str] = None
) -> List[dict]:
    """按知识点统计作答正确率（用于雷达图/进度条）；subject 非空时仅统计该科目。"""
    stmt = select(LearningRecord, Exercise).join(Exercise, Exercise.id == LearningRecord.exercise_id)
    cond = [LearningRecord.student_id == student_id]
    if subject:
        stmt = stmt.join(Course, Course.id == Exercise.course_id)
        cond.append(Course.subject == subject)
    stmt = stmt.where(*cond).order_by(desc(LearningRecord.created_at)).limit(limit)
    result = await db.execute(stmt)
    stats: dict = defaultdict(lambda: {"total": 0, "correct": 0})
    for rec, ex in result.all():
        for kp in ex.knowledge_points or []:
            stats[kp]["total"] += 1
            if rec.is_correct:
                stats[kp]["correct"] += 1
    items = [
        {
            "knowledge_point": kp,
            "total": v["total"],
            "correct": v["correct"],
            "mastery": round(v["correct"] / v["total"] * 100, 1) if v["total"] else 0,
        }
        for kp, v in stats.items()
        if v["total"] > 0
    ]
    items.sort(key=lambda x: x["mastery"])
    return items


async def _weak_points(db: AsyncSession, student_id: int, topn: int = 10) -> List[dict]:
    mastery = await _knowledge_mastery(db, student_id, limit=1500)
    weak = [m for m in mastery if m["mastery"] < 70][:topn]
    return weak or mastery[:topn]


async def _class_student_ids(db: AsyncSession, class_ids: List[int]) -> List[int]:
    if not class_ids:
        return []
    result = await db.execute(
        select(ClassStudent.student_id).where(ClassStudent.class_id.in_(class_ids))
    )
    return list(result.scalars().all())


async def _student_classes(db: AsyncSession, student_id: int) -> List[dict]:
    result = await db.execute(
        select(Class, User)
        .join(ClassStudent, ClassStudent.class_id == Class.id)
        .join(User, User.id == Class.teacher_id)
        .where(ClassStudent.student_id == student_id)
    )
    return [
        {
            "class_id": cls.id,
            "class_name": cls.name,
            "grade": cls.grade,
            "teacher_name": teacher.full_name or teacher.username,
        }
        for cls, teacher in result.all()
    ]


def _grade_order_expr(col):
    """年级顺序：一年级~九年级（含初一/初二/初三、高一/高二/高三），未知年级排最后。"""
    return case(
        (col == "一年级", 1), (col == "二年级", 2), (col == "三年级", 3),
        (col == "四年级", 4), (col == "五年级", 5), (col == "六年级", 6),
        (col == "七年级", 7), (col == "八年级", 8), (col == "九年级", 9),
        (col == "初一", 7), (col == "初二", 8), (col == "初三", 9),
        (col == "高一", 10), (col == "高二", 11), (col == "高三", 12),
        else_=99,
    )


def _normalize_teacher_no(raw: Optional[str]) -> str:
    """把教师工号统一为 t26xxxx 格式（小写 t + 26 + 4 位数字）。"""
    digits = "".join(ch for ch in (raw or "") if ch.isdigit())
    suffix = digits[-4:].rjust(4, "0")
    return f"t26{suffix}"


async def _next_teacher_no(db: AsyncSession) -> str:
    """自动生成下一个 t26xxxx 工号（取已有工号最大序号 + 1）。"""
    rows = (
        await db.execute(select(Teacher.employee_no).where(Teacher.employee_no.isnot(None)))
    ).scalars().all()
    nums = []
    for no in rows:
        digits = "".join(ch for ch in (no or "") if ch.isdigit())
        if digits:
            nums.append(int(digits[-4:]))
    return f"t26{max(nums) + 1 if nums else 1:04d}"


# ------------------------- 个人信息 -------------------------
@router.get("/me")
async def get_me(user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    student, teacher = await _load_profile(db, user)
    return ok(user_to_dict(user, student, teacher))


@router.put("/me")
async def update_me(
    payload: UserUpdate,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if payload.username is not None:
        user.username = payload.username
    if payload.full_name is not None:
        user.full_name = payload.full_name
    if payload.email is not None:
        user.email = payload.email or None
    if payload.avatar is not None:
        user.avatar = payload.avatar
    if payload.password:
        user.password_hash = hash_password(payload.password)
    await db.commit()
    await db.refresh(user)
    student, teacher = await _load_profile(db, user)
    return ok(user_to_dict(user, student, teacher), message="更新成功")


@router.post("/me/avatar")
async def upload_avatar(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    suffix = ("." + (file.filename or "").split(".")[-1]).lower()
    if suffix not in AVATAR_EXTS:
        return fail("仅支持 png/jpg/jpeg/gif/webp 图片")
    content = await file.read()
    if len(content) > 2 * 1024 * 1024:
        return fail("头像图片不能超过 2MB")
    import uuid

    avatar_dir = settings.upload_dir / "avatars"
    avatar_dir.mkdir(parents=True, exist_ok=True)
    name = f"{uuid.uuid4().hex}{suffix}"
    (avatar_dir / name).write_bytes(content)
    user.avatar = f"/uploads/avatars/{name}"
    await db.commit()
    return ok({"avatar": user.avatar}, message="头像已更新")

# ------------------------- 学生端：学习看板 -------------------------


@router.get("/me/courses")
async def my_courses(user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """课程列表（供学生选择科目/章节、AI 出题关联课程）。"""
    result = await db.execute(select(Course).order_by(Course.id))
    courses = result.scalars().all()
    return ok(
        [
            {
                "id": c.id,
                "name": c.name,
                "grade": c.grade,
                "subject": c.subject,
                "teacher_id": c.teacher_id,
                "chapter_tree": c.chapter_tree or [],
                "description": c.description,
            }
            for c in courses
        ]
    )

@router.get("/me/dashboard")
async def student_dashboard(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    student = await _get_student(db, user)
    now = datetime.now()
    week_ago = now - timedelta(days=7)

    # 答题统计（与 /me/stats 学情分析保持一致：按作答记录计数）
    total_answered = (
        await db.execute(
            select(func.count(LearningRecord.id)).where(
                LearningRecord.student_id == student.id
            )
        )
    ).scalar() or 0
    correct_count = (
        await db.execute(
            select(func.count(LearningRecord.id)).where(
                LearningRecord.student_id == student.id,
                LearningRecord.is_correct.is_(True),
            )
        )
    ).scalar() or 0
    total_exercises = (
        (await db.execute(select(func.count(Exercise.id)))).scalar() or 0
    )

    wrong_count = (
        await db.execute(
            select(func.count(WrongBook.id)).where(WrongBook.student_id == student.id)
        )
    ).scalar() or 0

    weekly_seconds = (
        await db.execute(
            select(func.coalesce(func.sum(LearningRecord.duration_seconds), 0)).where(
                LearningRecord.student_id == student.id,
                LearningRecord.created_at >= week_ago,
            )
        )
    ).scalar() or 0

    weak = await _weak_points(db, student.id, topn=10)
    mastery = await _knowledge_mastery(db, student.id)
    classes = await _student_classes(db, student.id)

    total_percent = round(total_answered / total_exercises * 100, 1) if total_exercises else 0
    return ok(
        {
            "user": user_to_dict(user, student),
            "classes": classes,
            "today_tasks": {
                "pending_exercises": max(total_exercises - total_answered, 0),
                "wrong_review": wrong_count,
            },
            "total_answered": total_answered,
            "correct_count": correct_count,
            "accuracy": round(correct_count / total_answered, 4) if total_answered else 0,
            "subject_count": len(mastery),
            "weekly_seconds": weekly_seconds,
            "total_percent": total_percent,
            "weak_points": weak,
        }
    )


@router.get("/me/today-schedule")
async def my_today_schedule(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """学生今日课程安排：自动取今天星期，按所在班级课表读取今天要上的课。"""
    student = await _get_student(db, user)
    class_ids = (
        await db.execute(select(ClassStudent.class_id).where(ClassStudent.student_id == student.id))
    ).scalars().all()
    if not class_ids:
        return ok([])
    today_weekday = date.today().isoweekday()
    rows = (
        await db.execute(
            select(ClassSchedule, Class, ClassScheduleTeacher)
            .join(Class, Class.id == ClassSchedule.class_id)
            .outerjoin(ClassScheduleTeacher, ClassScheduleTeacher.schedule_id == ClassSchedule.id)
            .where(
                ClassSchedule.class_id.in_(list(class_ids)),
                ClassSchedule.weekday == today_weekday,
                ClassSchedule.subject.isnot(None),
                ClassSchedule.subject != "",
            )
            .order_by(ClassSchedule.period)
        )
    ).all()
    # 上课老师姓名
    teacher_names: dict = {}
    for cs, _cls, link in rows:
        if link and link.teacher_user_id not in teacher_names:
            t_user = await db.get(User, link.teacher_user_id)
            teacher_names[link.teacher_user_id] = t_user.full_name or t_user.username if t_user else None
    return ok(
        [
            {
                "period": cs.period,
                "class_name": cls.name,
                "subject": cs.subject,
                "teacher_name": teacher_names.get(link.teacher_user_id) if link else None,
            }
            for cs, cls, link in rows
        ]
    )


@router.get("/me/week-schedule")
async def my_week_schedule(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """学生本周课表：按班级课表返回周一~周五全部课程，用于看板「本周」格子视图。"""
    student = await _get_student(db, user)
    class_ids = (
        await db.execute(select(ClassStudent.class_id).where(ClassStudent.student_id == student.id))
    ).scalars().all()
    if not class_ids:
        return ok([])
    rows = (
        await db.execute(
            select(ClassSchedule, Class, ClassScheduleTeacher)
            .join(Class, Class.id == ClassSchedule.class_id)
            .outerjoin(ClassScheduleTeacher, ClassScheduleTeacher.schedule_id == ClassSchedule.id)
            .where(
                ClassSchedule.class_id.in_(list(class_ids)),
                ClassSchedule.weekday.between(1, 5),
                ClassSchedule.subject.isnot(None),
                ClassSchedule.subject != "",
            )
            .order_by(ClassSchedule.weekday, ClassSchedule.period)
        )
    ).all()
    teacher_names: dict = {}
    for cs, _cls, link in rows:
        if link and link.teacher_user_id not in teacher_names:
            t_user = await db.get(User, link.teacher_user_id)
            teacher_names[link.teacher_user_id] = t_user.full_name or t_user.username if t_user else None
    return ok(
        [
            {
                "weekday": cs.weekday,
                "period": cs.period,
                "class_name": cls.name,
                "subject": cs.subject,
                "teacher_name": teacher_names.get(link.teacher_user_id) if link else None,
            }
            for cs, cls, link in rows
        ]
    )


# ------------------------- 学生端：学习时间线 -------------------------
@router.get("/me/timeline")
async def my_timeline(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    start: Optional[str] = Query(None, description="YYYY-MM-DD"),
    end: Optional[str] = Query(None, description="YYYY-MM-DD"),
    limit: int = Query(100, ge=1, le=300),
):
    student = await _get_student(db, user)
    items: List[dict] = []

    def in_range(dt):
        if dt is None:
            return True
        if start and dt.date() < date.fromisoformat(start):
            return False
        if end and dt.date() > date.fromisoformat(end):
            return False
        return True

    rows = (
        await db.execute(
            select(LearningRecord, Exercise)
            .join(Exercise, Exercise.id == LearningRecord.exercise_id)
            .where(LearningRecord.student_id == student.id)
            .order_by(desc(LearningRecord.created_at))
            .limit(300)
        )
    ).all()
    for rec, ex in rows:
        if not in_range(rec.created_at):
            continue
        items.append(
            {
                "time": rec.created_at.isoformat() if rec.created_at else None,
                "type": "exercise",
                "is_correct": rec.is_correct,
                "title": (ex.content or "")[:60],
                "detail": f"{'答对' if rec.is_correct else '答错'} · 用时 {rec.duration_seconds or 0}s",
                "exercise_id": rec.exercise_id,
            }
        )

    chats = (
        await db.execute(
            select(ChatMessage, ChatSession)
            .join(ChatSession, ChatSession.id == ChatMessage.session_id)
            .where(ChatMessage.role == "user", ChatSession.user_id == user.id)
            .order_by(desc(ChatMessage.created_at))
            .limit(200)
        )
    ).all()
    for msg, session in chats:
        if not in_range(msg.created_at):
            continue
        items.append(
            {
                "time": msg.created_at.isoformat() if msg.created_at else None,
                "type": "question",
                "title": (msg.content or "")[:60],
                "detail": f"会话：{session.title}",
                "session_id": session.id,
            }
        )

    wrongs = (
        await db.execute(
            select(WrongBook, Exercise)
            .join(Exercise, Exercise.id == WrongBook.exercise_id)
            .where(WrongBook.student_id == student.id)
            .order_by(desc(WrongBook.created_at))
            .limit(200)
        )
    ).all()
    for wb, ex in wrongs:
        if not in_range(wb.created_at):
            continue
        items.append(
            {
                "time": wb.created_at.isoformat() if wb.created_at else None,
                "type": "collect",
                "title": (ex.content or "")[:60],
                "detail": "收藏错题",
                "exercise_id": ex.id,
            }
        )

    items.sort(key=lambda x: x["time"] or "", reverse=True)
    return ok(items[:limit])


# ------------------------- 学生端：我的班级与任课教师 -------------------------
@router.get("/me/class-info")
async def my_class_info(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    student = await _get_student(db, user)
    result = await db.execute(
        select(Class, User, Teacher)
        .join(ClassStudent, ClassStudent.class_id == Class.id)
        .join(User, User.id == Class.teacher_id)
        .outerjoin(Teacher, Teacher.user_id == User.id)
        .where(ClassStudent.student_id == student.id)
    )
    classes = []
    for cls, teacher_user, teacher_ext in result.all():
        classes.append(
            {
                "class_id": cls.id,
                "class_name": cls.name,
                "grade": cls.grade,
                "teacher_name": teacher_user.full_name or teacher_user.username,
                "teacher_title": teacher_ext.title if teacher_ext else None,
                "teacher_department": teacher_ext.department if teacher_ext else None,
            }
        )
    return ok(classes)


# ------------------------- 学生端：学习记录 / 统计 -------------------------
@router.get("/me/records")
async def my_records(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    limit: int = Query(50, ge=1, le=200),
):
    student = await _get_student(db, user)
    result = await db.execute(
        select(LearningRecord, Exercise)
        .join(Exercise, LearningRecord.exercise_id == Exercise.id)
        .where(LearningRecord.student_id == student.id)
        .order_by(desc(LearningRecord.created_at))
        .limit(limit)
    )
    rows = result.all()
    return ok(
        [
            {
                "id": rec.id,
                "exercise_id": rec.exercise_id,
                "user_answer": rec.user_answer,
                "is_correct": rec.is_correct,
                "duration_seconds": rec.duration_seconds,
                "created_at": rec.created_at.isoformat() if rec.created_at else None,
                "exercise": exercise_to_dict(ex),
            }
            for rec, ex in rows
        ]
    )


@router.post("/me/records")
async def submit_record(
    payload: SubmitRecordRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    student = await _get_student(db, user)
    exercise = await db.get(Exercise, payload.exercise_id)
    if exercise is None:
        return fail("题目不存在")
    record = LearningRecord(
        student_id=student.id,
        exercise_id=payload.exercise_id,
        user_answer=payload.user_answer,
        is_correct=payload.is_correct,
        duration_seconds=payload.duration_seconds,
    )
    db.add(record)
    # 注：答错不再自动加入错题本，由学生点击“加入错题本”手动收藏
    await db.commit()
    await db.refresh(record)
    return ok({"id": record.id, "is_correct": record.is_correct}, message="已记录答题结果")


@router.get("/me/stats")
async def my_stats(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    subject: Optional[str] = Query(None, description="按科目过滤统计"),
):
    student = await _get_student(db, user)

    def _rec_stmt():
        stmt = (
            select(LearningRecord)
            .join(Exercise, Exercise.id == LearningRecord.exercise_id)
            .where(LearningRecord.student_id == student.id)
        )
        if subject:
            stmt = (
                stmt.join(Course, Course.id == Exercise.course_id)
                .where(Course.subject == subject)
            )
        return stmt

    total = (await db.execute(_rec_stmt().with_only_columns(func.count(LearningRecord.id)))).scalar() or 0
    correct = (
        await db.execute(
            _rec_stmt().with_only_columns(
                func.count(LearningRecord.id)
            ).where(LearningRecord.is_correct.is_(True))
        )
    ).scalar() or 0

    by_course_rows = (
        await db.execute(
            select(
                Course.name,
                func.count(LearningRecord.id),
                func.sum(case((LearningRecord.is_correct.is_(True), 1), else_=0)),
            )
            .join(Exercise, Exercise.id == LearningRecord.exercise_id)
            .join(Course, Course.id == Exercise.course_id)
            .where(LearningRecord.student_id == student.id)
            .group_by(Course.id, Course.name)
        )
    ).all()
    by_course = [
        {
            "course": name,
            "count": count,
            "correct": corr,
            "accuracy": round(corr / count, 4) if count else 0,
        }
        for name, count, corr in by_course_rows
    ]

    daily_stmt = (
        select(func.date(LearningRecord.created_at), func.count(LearningRecord.id))
        .join(Exercise, Exercise.id == LearningRecord.exercise_id)
        .where(LearningRecord.student_id == student.id)
    )
    if subject:
        daily_stmt = daily_stmt.join(Course, Course.id == Exercise.course_id).where(Course.subject == subject)
    daily_rows = (
        await db.execute(
            daily_stmt.group_by(func.date(LearningRecord.created_at))
            .order_by(func.date(LearningRecord.created_at))
        )
    ).all()
    daily = [{"date": str(d), "count": c} for d, c in daily_rows]

    mastery = await _knowledge_mastery(db, student.id, subject=subject)

    return ok(
        {
            "total_answered": total,
            "correct_count": correct,
            "accuracy": round(correct / total, 4) if total else 0,
            "by_course": by_course,
            "daily": daily,
            "knowledge_mastery": mastery,
        }
    )


def _week_monday(d: date) -> date:
    """返回 d 所在周的周一（周一为一周起点）。"""
    from app.services.activity_stats import week_monday
    return week_monday(d)


def _build_activity_series(daily: dict, period: str, start: date) -> list:
    """把按日期聚合的 {iso_date: count} 汇总为日/周/月序列。"""
    from app.services.activity_stats import build_activity_series
    return build_activity_series(daily, period, start)


@router.get("/me/activity")
async def my_activity(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    period: str = Query("day", pattern="^(day|week|month)$"),
    subject: Optional[str] = Query(None, description="按科目过滤活跃度"),
):
    """学生学习活跃度趋势（日/周/月）。

    仅统计该学生自身作答记录；date 窗口安全，避免全表扫描。
    """
    student = await _get_student(db, user)
    today = date.today()
    if period == "day":
        start = today - timedelta(days=13)
    elif period == "week":
        start = _week_monday(today - timedelta(weeks=7))
    else:  # month
        y, m = today.year, today.month
        for _ in range(5):
            m -= 1
            if m == 0:
                m = 12
                y -= 1
        start = date(y, m, 1)

    stmt = (
        select(func.date(LearningRecord.created_at), func.count(LearningRecord.id))
        .join(Exercise, Exercise.id == LearningRecord.exercise_id)
        .where(
            LearningRecord.student_id == student.id,
            func.date(LearningRecord.created_at) >= start,
        )
    )
    if subject:
        stmt = stmt.join(Course, Course.id == Exercise.course_id).where(Course.subject == subject)
    rows = (await db.execute(stmt.group_by(func.date(LearningRecord.created_at)))).all()
    daily = {str(d): int(c) for d, c in rows}
    return ok({"period": period, "items": _build_activity_series(daily, period, start)})


# ------------------------- 错题本 -------------------------
@router.get("/me/wrong-book")
async def my_wrong_book(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    knowledge_point: Optional[str] = Query(None),
    subject: Optional[str] = Query(None),
):
    student = await _get_student(db, user)
    result = await db.execute(
        select(WrongBook, Exercise, Course)
        .join(Exercise, WrongBook.exercise_id == Exercise.id)
        .join(Course, Course.id == Exercise.course_id, isouter=True)
        .where(WrongBook.student_id == student.id)
        .order_by(desc(WrongBook.created_at))
    )
    rows = result.all()
    items = []
    for wb, ex, course in rows:
        kps = ex.knowledge_points or []
        if knowledge_point and knowledge_point not in kps:
            continue
        if subject and (course is None or course.subject != subject):
            continue
        items.append(
            {
                "id": wb.id,
                "exercise_id": wb.exercise_id,
                "reason": wb.reason,
                "created_at": wb.created_at.isoformat() if wb.created_at else None,
                "subject": course.subject if course else "",
                "exercise": exercise_to_dict(ex),
            }
        )
    return ok(items)


@router.post("/me/wrong-book/batch-delete")
async def batch_delete_wrong_book(
    ids: list[int],
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    student = await _get_student(db, user)
    result = await db.execute(
        delete(WrongBook).where(
            WrongBook.student_id == student.id,
            WrongBook.id.in_(ids),
        )
    )
    await db.commit()
    return ok(None, message=f"已删除 {result.rowcount} 条错题")


@router.post("/me/wrong-book")
async def add_wrong_book(
    payload: WrongBookAddRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """把题目显式加入错题本（错题收藏）。"""
    student = await _get_student(db, user)
    if await db.get(Exercise, payload.exercise_id) is None:
        return fail("题目不存在")
    existing = await db.execute(
        select(WrongBook).where(
            WrongBook.student_id == student.id, WrongBook.exercise_id == payload.exercise_id
        )
    )
    if existing.scalar_one_or_none() is not None:
        return fail("该题已在错题本中")
    db.add(WrongBook(student_id=student.id, exercise_id=payload.exercise_id, reason=payload.reason))
    await db.commit()
    return ok(None, message="已加入错题本")


@router.delete("/me/wrong-book/{wrong_id}")
async def delete_wrong_book(
    wrong_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    student = await _get_student(db, user)
    wb = await db.get(WrongBook, wrong_id)
    if wb is None or wb.student_id != student.id:
        raise HTTPException(status_code=404, detail="错题记录不存在")
    await db.delete(wb)
    await db.commit()
    return ok(None, message="已删除")


# ------------------------- 练习 -------------------------
@router.get("/me/exercises")
async def get_exercises(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    course_id: Optional[int] = Query(None),
    chapter: Optional[str] = Query(None),
    difficulty: Optional[str] = Query(None),
    knowledge_point: Optional[str] = Query(None),
    limit: int = Query(10, ge=1, le=50),
):
    conditions = []
    if course_id:
        conditions.append(Exercise.course_id == course_id)
    if chapter:
        conditions.append(Exercise.chapter == chapter)
    if difficulty:
        conditions.append(Exercise.difficulty == difficulty)
    if knowledge_point:
        conditions.append(Exercise.knowledge_points.contains([knowledge_point]))
    stmt = select(Exercise)
    if conditions:
        stmt = stmt.where(*conditions)
    stmt = stmt.order_by(func.rand()).limit(limit)
    result = await db.execute(stmt)
    exercises = result.scalars().all()
    return ok([exercise_to_dict(e) for e in exercises])


@router.get("/me/similar-exercises")
async def similar_exercises(
    exercise_id: int = Query(..., ge=1, description="原题ID"),
    limit: int = Query(3, ge=1, le=10, description="返回条数"),
    allow_generate: bool = Query(True, description="题库不足时是否允许 AI 变式兜底"),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """相似例题推荐（P0 门限/硬过滤/回退 + P1 语义打标 + P2 AI 变式兜底）。

    - P0 硬过滤：题型必须一致；可判定科目必须一致；
      原题有知识点时候选必须至少命中一个语义知识点（剔除《蒹葭》类无关混入）；
      归一化相似度低于门限不推；质量不足时 kp →（无标签时）章节/课程 逐级回退，宁缺毋滥。
    - P1：knowledge_points 规范化 + 学科同义词/上下位别名扩展，语义匹配而非字符串相等。
    - P2：高质量候选不足且允许时，调用 LLM 生成同知识点/同题型/同难度的变式题（内存返回，不写库）。
    不修改任何数据库结构/数据模型/存储。
    """
    orig = await db.get(Exercise, exercise_id)
    if orig is None:
        return fail("题目不存在")

    rows = (await db.execute(select(Exercise))).scalars().all()

    # 批量加载课程科目，避免逐条查询
    course_ids = {ex.course_id for ex in rows if ex.course_id} | {orig.course_id}
    course_map = {}
    if course_ids:
        course_rows = (
            await db.execute(select(Course).where(Course.id.in_(course_ids)))
        ).scalars().all()
        course_map = {c.id: c for c in course_rows}

    def _feature(ex: Exercise) -> ExerciseFeature:
        course = course_map.get(ex.course_id)
        return ExerciseFeature(
            id=ex.id,
            type=ex.type,
            difficulty=ex.difficulty,
            subject=(course.subject if course else None) or "",
            chapter=ex.chapter or "",
            course_id=ex.course_id,
            kps=ex.knowledge_points or [],
        )

    picked, level, insufficient = select_similar(
        _feature(orig), [_feature(c) for c in rows], limit=limit, threshold=0.6
    )
    by_id = {c.id: c for c in rows}
    results = [dict(exercise_to_dict(by_id[c.id])) for _, c in picked]

    # P2：库中相似题不足 limit 道时，用 AI 生成同构变式补齐（内存返回，不落库）
    need = limit - len(results)
    generated = []
    if need > 0 and allow_generate:
        try:
            generated = await _generate_similar_fallback(orig, need)
        except Exception:  # noqa: BLE001  AI 兜底失败不影响主流程
            logger.warning("相似题 AI 兜底生成失败，exercise_id=%s", exercise_id, exc_info=True)
            generated = []

    seen = {_norm(exercise_to_dict(by_id[c.id])["content"]) for _, c in picked}
    fused = list(results)
    gen_quality = []
    for g in generated:
        key = _norm(g.get("content") or "")
        if not key or key in seen:
            continue
        seen.add(key)
        g["generated"] = True
        fused.append(g)
        gen_quality.append(_gen_quality(g))

    return ok(
        {
            "items": fused,
            "level": level,
            "insufficient": len(fused) < limit,
            "generated_count": len(gen_quality),
            "quality": gen_quality,
        },
        message="相似例题推荐完成",
    )


def _norm(text: str) -> str:
    return "".join((text or "").strip().split()).lower()


def _gen_quality(g: dict) -> dict:
    """生成结果的 AI 质量评估标记（供前端/测试观察，不含 DB 持久化）。"""
    content_ok = bool((g.get("content") or "").strip())
    answer_ok = bool((g.get("answer") or "").strip())
    kp_ok = bool((g.get("knowledge_point") or "").strip())
    ai_quality = (g.get("quality") or "low").lower()
    if content_ok and answer_ok and kp_ok:
        overall = ai_quality if ai_quality in ("high", "medium") else "low"
    else:
        overall = "low"
    return {
        "content_ok": content_ok,
        "answer_ok": answer_ok,
        "knowledge_point_ok": kp_ok,
        "ai_quality": ai_quality,
        "overall": overall,
    }


async def _generate_similar_fallback(orig: Exercise, limit: int) -> list:
    """基于原题生成同构变式题；仅返回内存数据，绝不写入数据库。

    P2 质量门槛：content/answer/knowledge_point 齐全、ai quality 非 low、
    且生成题与原知识点语义一致（同义词收敛后必须命中）才放行。
    """
    original = json.dumps(exercise_to_dict(orig), ensure_ascii=False, default=str)
    prompt = SIMILAR_EXERCISE_PROMPT.format(original=original)
    out = []
    seen = {_norm(orig.content)}
    for _ in range(limit * 3):
        try:
            text = await llm_client.chat(
                [{"role": "system", "content": prompt}], temperature=0.7, max_tokens=1024
            )
            data = parse_json(text)
        except Exception:
            logger.warning("单次 LLM 生成相似题失败，跳过本次重试", exc_info=True)
            continue
        data = data or {}
        if not _passes_gen_gate(data, orig, seen):
            continue
        out.append(_coerce_generated(data, orig))
        if len(out) >= limit:
            break
    return out


def _passes_gen_gate(data: dict, orig: Exercise, seen: set) -> bool:
    content = data.get("content") or ""
    answer = data.get("answer") or ""
    kp = data.get("knowledge_point") or ""
    if not (content and answer and kp):
        return False
    if _norm(content) in seen:
        return False
    if data.get("quality") == "low":
        return False
    # 知识点语义校验：仅当原题知识点在 KP_SYNONYMS 中有同义词映射时严格匹配
    orig_kps = orig.knowledge_points or []
    if orig_kps:
        orig_canon = canonical_kps(orig_kps)
        has_synonyms = any(c in KP_SYNONYMS for c in orig_canon)
        if has_synonyms:
            orig_exp = expanded_kps(orig_kps)
            if orig_exp and not (orig_exp & expanded_kps([kp])):
                return False
    seen.add(_norm(content))
    return True


def _coerce_generated(data: dict, orig: Exercise) -> dict:
    """将模型输出统一为前端可用结构（携带原题上下文，id 置空标识为生成题）。"""
    return {
        "id": None,
        "course_id": orig.course_id,
        "chapter": orig.chapter,
        "type": orig.type,
        "content": data.get("content", ""),
        "options": data.get("options") or [],
        "answer": data.get("answer", ""),
        "analysis": data.get("analysis", ""),
        "difficulty": data.get("difficulty") or orig.difficulty,
        "knowledge_points": [data["knowledge_point"]] if data.get("knowledge_point") else [],
        "generated": True,
    }


# ------------------------- 教师端：工作台概览 -------------------------
@router.get("/me/teacher-dashboard")
async def teacher_dashboard(
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    teacher = await _get_teacher(db, user) if user.role == "teacher" else None
    if user.role == "teacher":
        result = await db.execute(select(Class).where(Class.teacher_id == user.id))
    else:
        result = await db.execute(select(Class))
    classes = result.scalars().all()
    class_ids = [c.id for c in classes]
    student_ids = await _class_student_ids(db, class_ids)

    total_records = correct = 0
    if student_ids:
        total_records = (
            await db.execute(
                select(func.count(LearningRecord.id)).where(
                    LearningRecord.student_id.in_(student_ids)
                )
            )
        ).scalar() or 0
        correct = (
            await db.execute(
                select(func.count(LearningRecord.id)).where(
                    LearningRecord.student_id.in_(student_ids),
                    LearningRecord.is_correct.is_(True),
                )
            )
        ).scalar() or 0

    week_ago = datetime.now() - timedelta(days=7)
    daily_rows = []
    if student_ids:
        daily_rows = (
            await db.execute(
                select(func.date(LearningRecord.created_at), func.count(LearningRecord.id))
                .where(
                    LearningRecord.student_id.in_(student_ids),
                    LearningRecord.created_at >= week_ago,
                )
                .group_by(func.date(LearningRecord.created_at))
                .order_by(func.date(LearningRecord.created_at))
            )
        ).all()

    # 今日安排：自动取今天星期，在课表中查询自己上课的课
    today_weekday = date.today().isoweekday()
    schedule_rows = (
        await db.execute(
            select(ClassSchedule, Class)
            .join(ClassScheduleTeacher, ClassScheduleTeacher.schedule_id == ClassSchedule.id)
            .join(Class, Class.id == ClassSchedule.class_id)
            .where(
                ClassSchedule.weekday == today_weekday,
                ClassSchedule.subject.isnot(None),
                ClassSchedule.subject != "",
                ClassScheduleTeacher.teacher_user_id == user.id,
            )
            .order_by(ClassSchedule.period)
        )
    ).all()
    today_schedule = [
        {
            "period": cs.period,
            "class_id": cls.id,
            "class_name": cls.name,
            "subject": cs.subject,
        }
        for cs, cls in schedule_rows
    ]

    return ok(
        {
            "user": user_to_dict(user, None, teacher),
            "class_count": len(classes),
            "student_count": len(student_ids),
            "total_answered": total_records,
            "accuracy": round(correct / total_records, 4) if total_records else 0,
            "daily_activity": [{"date": str(d), "count": c} for d, c in daily_rows],
            "today_schedule": today_schedule,
        }
    )


# ------------------------- 教师端：班级列表与详情 -------------------------
@router.get("/me/classes")
async def my_classes(
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    if user.role == "teacher":
        result = await db.execute(select(Class).where(Class.teacher_id == user.id).order_by(Class.id))
    else:
        result = await db.execute(select(Class).order_by(Class.id))
    classes = result.scalars().all()
    items = []
    for cls in classes:
        count = (
            await db.execute(
                select(func.count(ClassStudent.id)).where(ClassStudent.class_id == cls.id)
            )
        ).scalar() or 0
        items.append(
            {
                "id": cls.id,
                "name": cls.name,
                "class_no": cls.class_no,
                "grade": cls.grade,
                "description": cls.description,
                "student_count": count,
            }
        )
    return ok(items)


@router.get("/me/classes/{class_id}")
async def class_detail(
    class_id: int,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    cls = await db.get(Class, class_id)
    if cls is None:
        raise HTTPException(status_code=404, detail="班级不存在")
    if user.role == "teacher" and cls.teacher_id != user.id:
        raise HTTPException(status_code=403, detail="无权查看该班级")

    result = await db.execute(
        select(Student, User, ClassStudent)
        .join(ClassStudent, ClassStudent.student_id == Student.id)
        .join(User, User.id == Student.user_id)
        .where(ClassStudent.class_id == class_id)
    )
    rows = result.all()
    student_ids = [s.id for s, _, _ in rows]

    students = []
    for s, u, _ in rows:
        s_total = (
            await db.execute(
                select(func.count(LearningRecord.id)).where(LearningRecord.student_id == s.id)
            )
        ).scalar() or 0
        s_correct = (
            await db.execute(
                select(func.count(LearningRecord.id)).where(
                    LearningRecord.student_id == s.id, LearningRecord.is_correct.is_(True)
                )
            )
        ).scalar() or 0
        last = (
            await db.execute(
                select(LearningRecord.created_at)
                .where(LearningRecord.student_id == s.id)
                .order_by(desc(LearningRecord.created_at))
                .limit(1)
            )
        ).scalar()
        students.append(
            {
                "student_id": s.id,
                "user_id": u.id,
                "username": u.username,
                "full_name": u.full_name or u.username,
                "student_no": s.student_no,
                "grade": s.grade,
                "answered": s_total,
                "correct": s_correct,
                "accuracy": round(s_correct / s_total, 4) if s_total else 0,
                "last_active": last.isoformat() if last else None,
            }
        )

    total = sum(x["answered"] for x in students)
    correct = sum(x["correct"] for x in students)
    weak_counter: Counter = Counter()
    if student_ids:
        recs = (
            await db.execute(
                select(LearningRecord, Exercise)
                .join(Exercise, Exercise.id == LearningRecord.exercise_id)
                .where(
                    LearningRecord.student_id.in_(student_ids),
                    LearningRecord.is_correct.is_(False),
                )
                .limit(1000)
            )
        ).all()
        for _, ex in recs:
            for kp in ex.knowledge_points or []:
                weak_counter[kp] += 1
    weak_top = [{"knowledge_point": k, "count": v} for k, v in weak_counter.most_common(5)]

    teacher_user = await db.get(User, cls.teacher_id)
    return ok(
        {
            "id": cls.id,
            "name": cls.name,
            "class_no": cls.class_no,
            "grade": cls.grade,
            "description": cls.description,
            "teacher_name": (teacher_user.full_name or teacher_user.username) if teacher_user else None,
            "student_count": len(students),
            "total_answered": total,
            "accuracy": round(correct / total, 4) if total else 0,
            "weak_top": weak_top,
            "students": students,
        }
    )


@router.delete("/me/classes/{class_id}/students/{student_id}")
async def remove_class_student(
    class_id: int,
    student_id: int,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    cls = await db.get(Class, class_id)
    if cls is None:
        raise HTTPException(status_code=404, detail="班级不存在")
    if user.role == "teacher" and cls.teacher_id != user.id:
        raise HTTPException(status_code=403, detail="无权操作该班级")
    link = (
        await db.execute(
            select(ClassStudent).where(
                ClassStudent.class_id == class_id,
                ClassStudent.student_id == student_id,
            )
        )
    ).scalar_one_or_none()
    if link is None:
        return fail("该学生不在此班级")
    await db.delete(link)
    await db.commit()
    return ok(None, message="已将该学生移出班级")

# ------------------------- 班级学习报告 -------------------------
@router.get("/me/classes/{class_id}/report")
async def class_report(
    class_id: int,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    detail = (await class_detail(class_id, user, db))["data"]
    return ok(detail, message="班级学习报告生成成功")


@router.get("/me/classes/{class_id}/report/export")
async def class_report_export(
    class_id: int,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    """导出班级学习报告为 Word (.docx)。"""
    detail = (await class_detail(class_id, user, db))["data"]
    doc = docx.Document()
    doc.add_heading(f"{detail['name']} 学习报告", 0)
    doc.add_paragraph(f"年级：{detail.get('grade') or '-'}    学科：{detail.get('subject') or '-'}")
    doc.add_paragraph(f"班主任/任课教师：{detail.get('teacher_name') or '-'}")
    doc.add_paragraph(
        f"学生人数：{detail['student_count']}    累计答题：{detail['total_answered']}"
        f"    平均正确率：{round(detail['accuracy'] * 100, 1)}%"
    )
    doc.add_heading("班级薄弱知识点 TOP5", level=1)
    for w in detail.get("weak_top") or []:
        doc.add_paragraph(f"{w['knowledge_point']}（错题 {w['count']} 次）", style="List Bullet")
    doc.add_heading("学生明细", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["姓名", "学号", "答题数", "答对数", "正确率"]):
        hdr[i].text = h
    for s in detail.get("students") or []:
        row = table.add_row().cells
        row[0].text = s["full_name"]
        row[1].text = s["student_no"] or "-"
        row[2].text = str(s["answered"])
        row[3].text = str(s["correct"])
        row[4].text = f"{round(s['accuracy'] * 100, 1)}%"
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = quote(f"{detail['name']}_学习报告.docx")
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


# ------------------------- 教师端：学生学情画像 -------------------------
@router.get("/me/students/{student_id}/profile")
async def student_profile(
    student_id: int,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    student = await db.get(Student, student_id)
    if student is None:
        raise HTTPException(status_code=404, detail="学生不存在")
    s_user = await db.get(User, student.user_id)
    if user.role == "teacher":
        t_class_ids = (
            await db.execute(select(Class.id).where(Class.teacher_id == user.id))
        ).scalars().all()
        link = await db.execute(
            select(ClassStudent).where(
                ClassStudent.student_id == student_id,
                ClassStudent.class_id.in_(list(t_class_ids)),
            )
        )
        if link.scalar_one_or_none() is None:
            raise HTTPException(status_code=403, detail="该学生不在您所带的班级中")

    classes = await _student_classes(db, student_id)
    mastery = await _knowledge_mastery(db, student_id, limit=1500)
    wrongs = (
        await db.execute(
            select(WrongBook, Exercise)
            .join(Exercise, Exercise.id == WrongBook.exercise_id)
            .where(WrongBook.student_id == student_id)
            .order_by(desc(WrongBook.created_at))
            .limit(100)
        )
    ).all()
    wrong_history = [
        {
            "id": wb.id,
            "exercise_id": wb.exercise_id,
            "reason": wb.reason,
            "created_at": wb.created_at.isoformat() if wb.created_at else None,
            "exercise": exercise_to_dict(ex),
        }
        for wb, ex in wrongs
    ]

    thirty_days = datetime.now() - timedelta(days=30)
    activity = (
        await db.execute(
            select(func.date(LearningRecord.created_at), func.count(LearningRecord.id))
            .where(LearningRecord.student_id == student_id, LearningRecord.created_at >= thirty_days)
            .group_by(func.date(LearningRecord.created_at))
            .order_by(func.date(LearningRecord.created_at))
        )
    ).all()
    behavior = (
        await db.execute(
            select(LearningRecord, Exercise)
            .join(Exercise, Exercise.id == LearningRecord.exercise_id)
            .where(LearningRecord.student_id == student_id)
            .order_by(desc(LearningRecord.created_at))
            .limit(30)
        )
    ).all()
    behavior_records = [
        {
            "time": rec.created_at.isoformat() if rec.created_at else None,
            "content": (ex.content or "")[:60],
            "is_correct": rec.is_correct,
            "exercise_id": rec.exercise_id,
        }
        for rec, ex in behavior
    ]
    messages = (
        await db.execute(
            select(StudentMessage, User)
            .join(User, User.id == StudentMessage.teacher_id)
            .where(StudentMessage.student_id == student_id)
            .order_by(desc(StudentMessage.created_at))
            .limit(50)
        )
    ).all()
    teacher_msgs = [
        {
            "id": m.id,
            "teacher_name": t.full_name or t.username,
            "content": m.content,
            "created_at": m.created_at.isoformat() if m.created_at else None,
        }
        for m, t in messages
    ]

    total = sum(a[1] for a in activity)
    return ok(
        {
            "student": {
                "user_id": s_user.id,
                "username": s_user.username,
                "full_name": s_user.full_name or s_user.username,
                "avatar": s_user.avatar,
                "student_no": student.student_no,
                "grade": student.grade,
            },
            "classes": classes,
            "knowledge_mastery": mastery,
            "wrong_history": wrong_history,
            "activity_30d": [{"date": str(d), "count": c} for d, c in activity],
            "behavior_records": behavior_records,
            "messages": teacher_msgs,
            "total_answered_30d": total,
        }
    )


@router.post("/me/students/{student_id}/message")
async def send_student_message(
    student_id: int,
    payload: StudentMessageCreate,
    user: User = Depends(require_roles("admin", "teacher")),
    db: AsyncSession = Depends(get_db),
):
    student = await db.get(Student, student_id)
    if student is None:
        raise HTTPException(status_code=404, detail="学生不存在")
    if user.role == "teacher":
        t_class_ids = (
            await db.execute(select(Class.id).where(Class.teacher_id == user.id))
        ).scalars().all()
        link = await db.execute(
            select(ClassStudent).where(
                ClassStudent.student_id == student_id,
                ClassStudent.class_id.in_(list(t_class_ids)),
            )
        )
        if link.scalar_one_or_none() is None:
            raise HTTPException(status_code=403, detail="该学生不在您所带的班级中")
    db.add(StudentMessage(student_id=student_id, teacher_id=user.id, content=payload.content))
    await db.commit()
    return ok(None, message="留言已发送")


# ------------------------- 班级学情概览（教师首页快捷） -------------------------
@router.get("")
async def list_users(
    user: User = Depends(require_roles("admin")),
    db: AsyncSession = Depends(get_db),
    role: Optional[str] = Query(None),
    keyword: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    size: int = Query(20, ge=1, le=100),
):
    stmt = select(User)
    if role:
        stmt = stmt.where(User.role == role)
    if keyword:
        stmt = stmt.where(User.username.ilike(f"%{keyword}%") | User.full_name.ilike(f"%{keyword}%"))
    # 排序：教师按工号升序；学生按年级顺序（七年级→八年级→九年级）再按学号升序
    if role == "teacher":
        stmt = stmt.join(Teacher, Teacher.user_id == User.id).order_by(Teacher.employee_no)
    elif role == "student":
        stmt = stmt.join(Student, Student.user_id == User.id).order_by(_grade_order_expr(Student.grade), Student.student_no)
    else:
        stmt = stmt.order_by(User.id)
    total = (
        await db.execute(select(func.count()).select_from(stmt.subquery()))
    ).scalar() or 0
    result = await db.execute(stmt.offset((page - 1) * size).limit(size))
    users = result.scalars().all()
    items = []
    for u in users:
        student = teacher = None
        class_names = []
        if u.role == "student":
            student = (
                await db.execute(select(Student).where(Student.user_id == u.id))
            ).scalar_one_or_none()
            # 学生所属班级（管理员列表显示用）
            if student is not None:
                class_names = (
                    await db.execute(
                        select(Class.name)
                        .join(ClassStudent, ClassStudent.class_id == Class.id)
                        .where(ClassStudent.student_id == student.id)
                    )
                ).scalars().all()
        elif u.role == "teacher":
            teacher = (
                await db.execute(select(Teacher).where(Teacher.user_id == u.id))
            ).scalar_one_or_none()
        item = user_to_dict(u, student, teacher)
        if student is not None:
            item["student_id"] = student.id
            item["class_name"] = "、".join(class_names)
        items.append(item)
    return ok({"total": total, "items": items})


@router.post("")
async def create_user(
    payload: AdminUserCreate,
    user: User = Depends(require_roles("admin")),
    db: AsyncSession = Depends(get_db),
):
    # 教师：负责科目必须为系统规定科目，且一位老师只负责一门学科
    if payload.role == "teacher" and payload.subjects:
        bad = [s for s in payload.subjects if not is_valid_subject(s)]
        if bad:
            return fail(f"存在非系统规定学科：{', '.join(bad)}")
        if len(payload.subjects) > 1:
            return fail("一位老师只能负责一门学科")
    # 教师：工号统一为 t26xxxx 格式（未填写则自动生成），且唯一
    employee_no = None
    if payload.role == "teacher":
        employee_no = _normalize_teacher_no(payload.employee_no) if payload.employee_no else await _next_teacher_no(db)
        dup = await db.execute(select(Teacher).where(Teacher.employee_no == employee_no))
        if dup.scalar_one_or_none() is not None:
            return fail("该工号已被其他教师使用")
    # 学生：学号唯一
    if payload.role == "student" and payload.student_no:
        dup = await db.execute(select(Student).where(Student.student_no == payload.student_no))
        if dup.scalar_one_or_none() is not None:
            return fail("该学号已被其他学生使用")
    new_user = User(
        username=payload.username,
        password_hash=hash_password(payload.password),
        full_name=payload.full_name,
        email=payload.email or None,
        role=payload.role,
        is_active=True,
    )
    db.add(new_user)
    await db.flush()
    if new_user.role == "teacher":
        db.add(
            Teacher(
                user_id=new_user.id,
                employee_no=employee_no,
                subjects=payload.subjects or [],
                title=payload.title,
                department=payload.department,
            )
        )
    elif new_user.role == "student":
        db.add(
            Student(
                user_id=new_user.id,
                student_no=payload.student_no or f"S{new_user.id:06d}",
                grade=payload.grade,
            )
        )
    await db.commit()
    await db.refresh(new_user)
    student, teacher = await _load_profile(db, new_user)
    data = user_to_dict(new_user, student, teacher)
    if student is not None:
        data["student_id"] = student.id
    return ok(data, message="创建成功")


@router.put("/{user_id}")
async def update_user(
    user_id: int,
    payload: AdminUserUpdate,
    user: User = Depends(require_roles("admin")),
    db: AsyncSession = Depends(get_db),
):
    target = await db.get(User, user_id)
    if target is None:
        raise HTTPException(status_code=404, detail="用户不存在")
    if payload.role is not None and payload.role != target.role:
        target.role = payload.role
        await db.flush()
        if payload.role == "teacher":
            if (
                await db.execute(select(Teacher).where(Teacher.user_id == user_id))
            ).scalar_one_or_none() is None:
                db.add(Teacher(user_id=user_id))
        elif payload.role == "student":
            if (
                await db.execute(select(Student).where(Student.user_id == user_id))
            ).scalar_one_or_none() is None:
                db.add(Student(user_id=user_id, student_no=f"S{user_id:06d}"))
    if payload.is_active is not None:
        target.is_active = payload.is_active
    if payload.full_name is not None:
        target.full_name = payload.full_name
    if payload.email is not None:
        target.email = payload.email or None
    if payload.new_password:
        target.password_hash = hash_password(payload.new_password)
    if target.role == "student":
        student = (
            await db.execute(select(Student).where(Student.user_id == user_id))
        ).scalar_one_or_none()
        if student:
            if payload.grade is not None:
                student.grade = payload.grade
    if target.role == "teacher":
        teacher = (
            await db.execute(select(Teacher).where(Teacher.user_id == user_id))
        ).scalar_one_or_none()
        if teacher:
            if payload.subjects is not None:
                bad = [s for s in payload.subjects if not is_valid_subject(s)]
                if bad:
                    return fail(f"存在非系统规定学科：{', '.join(bad)}")
                if len(payload.subjects) > 1:
                    return fail("一位老师只能负责一门学科")
                teacher.subjects = payload.subjects
            if payload.employee_no is not None:
                employee_no = _normalize_teacher_no(payload.employee_no)
                dup = await db.execute(
                    select(Teacher).where(
                        Teacher.employee_no == employee_no,
                        Teacher.user_id != user_id,
                    )
                )
                if dup.scalar_one_or_none() is not None:
                    return fail("该工号已被其他教师使用")
                teacher.employee_no = employee_no
            if payload.title is not None:
                teacher.title = payload.title
            if payload.department is not None:
                teacher.department = payload.department
    await db.commit()
    await db.refresh(target)
    student, teacher = await _load_profile(db, target)
    return ok(user_to_dict(target, student, teacher), message="更新成功")


@router.delete("/{user_id}")
async def delete_user(
    user_id: int,
    user: User = Depends(require_roles("admin")),
    db: AsyncSession = Depends(get_db),
):
    """删除用户：同步清理其所有关联数据，保证各端数据一致。"""
    if user_id == user.id:
        return fail("不能删除当前登录账号")
    target = await db.get(User, user_id)
    if target is None:
        raise HTTPException(status_code=404, detail="用户不存在")
    if target.role == "teacher":
        # 教师名下还有班级时不允许删除（避免班级失联）
        classes = (
            await db.execute(select(Class).where(Class.teacher_id == user_id))
        ).scalars().all()
        if classes:
            names = "、".join(c.name for c in classes)
            return fail(f"该教师名下还有班级（{names}），请先调整班级任课教师后再删除")
        await db.execute(
            delete(ClassScheduleTeacher).where(
                ClassScheduleTeacher.teacher_user_id == user_id
            )
        )
        await db.execute(delete(LessonPlan).where(LessonPlan.teacher_id == user_id))
        await db.execute(delete(StudentMessage).where(StudentMessage.teacher_id == user_id))
    elif target.role == "student":
        student = (
            await db.execute(select(Student).where(Student.user_id == user_id))
        ).scalar_one_or_none()
        if student is not None:
            await db.execute(delete(ClassStudent).where(ClassStudent.student_id == student.id))
            await db.execute(delete(LearningRecord).where(LearningRecord.student_id == student.id))
            await db.execute(delete(WrongBook).where(WrongBook.student_id == student.id))
            await db.execute(delete(StudentMessage).where(StudentMessage.student_id == student.id))
            await db.execute(delete(AgentTask).where(AgentTask.student_id == student.id))
    # 清理该用户的问答会话（含消息）
    sessions = (
        await db.execute(select(ChatSession.id).where(ChatSession.user_id == user_id))
    ).scalars().all()
    if sessions:
        await db.execute(delete(ChatMessage).where(ChatMessage.session_id.in_(list(sessions))))
        await db.execute(delete(ChatSession).where(ChatSession.user_id == user_id))
    await db.delete(target)
    await db.commit()
    return ok(None, message="删除成功")
