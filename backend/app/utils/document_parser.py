"""文档解析：支持 .pdf（pdfplumber）、.docx（python-docx）、.txt、.md。"""
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class PageContent:
    """一页（或一份文档）的文本内容。"""
    page: int
    text: str


def parse_pdf(path: str) -> List[PageContent]:
    """使用 pdfplumber 逐页提取文本，保留页码。"""
    import pdfplumber

    pages: List[PageContent] = []
    with pdfplumber.open(path) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(PageContent(page=idx, text=text))
    return pages


def parse_docx(path: str) -> List[PageContent]:
    """使用 python-docx 提取段落与表格文本。"""
    import docx

    doc = docx.Document(path)
    parts: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    text = "\n".join(parts)
    return [PageContent(page=1, text=text)] if text else []


def parse_txt(path: str) -> List[PageContent]:
    raw = Path(path).read_text(encoding="utf-8", errors="ignore")
    return [PageContent(page=1, text=raw)] if raw.strip() else []


def parse_document(path: str, filename: str) -> List[PageContent]:
    """按扩展名分发解析，返回带页码的文本列表。"""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix in (".txt", ".md"):
        return parse_txt(path)
    raise ValueError(f"不支持的文件类型: {suffix or '未知'}")